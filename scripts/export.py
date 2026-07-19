#!/usr/bin/env python3
"""
豆包聊天记录导出工具
支持 Markdown、Word、JSON、JSONL、HTML 格式导出。
"""

import json
import os
import sqlite3
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import List
from html import escape

import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_config():
    config_path = Path(__file__).parent / "config.yaml"
    if not config_path.exists():
        config_path = Path(__file__).parent.parent / "config.yaml"
    with open(config_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_all_chats(backup_dir):
    chats = []
    json_dir = backup_dir / "json"
    if not json_dir.exists():
        return chats
    for f in sorted(json_dir.glob("*.json")):
        try:
            with open(f, "r", encoding="utf-8") as fp:
                chats.append(json.load(fp))
        except:
            pass
    return chats


def get_chat_stats(chats):
    total_messages = sum(len(c.get("messages", [])) for c in chats)
    total_words = sum(sum(len(m.get("content", "")) for m in c.get("messages", [])) for c in chats)
    user_words = sum(sum(len(m.get("content", "")) for m in c.get("messages", []) if m.get("role") == "user") for c in chats)
    return {
        "total_chats": len(chats),
        "total_messages": total_messages,
        "total_words": total_words,
        "user_words": user_words,
        "assistant_words": total_words - user_words,
    }


# ========== Markdown 导出 ==========

def export_markdown(chats, output_path):
    lines = ["# 豆包聊天记录备份\n", f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n", f"共 {len(chats)} 个对话\n", "---\n"]
    for chat in chats:
        title = chat.get("title", "未命名对话")
        messages = chat.get("messages", [])
        lines.append(f"\n## {title}\n")
        lines.append(f"- **消息数**: {len(messages)}\n")
        for msg in messages:
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            role_label = "🤖 Assistant" if role == "assistant" else "👤 User" if role == "user" else f"❓ {role}"
            lines.append(f"### {role_label}\n")
            lines.append(f"{content}\n")
            lines.append("---\n")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[✓] Markdown: {output_path}")


# ========== Word 导出 ==========

def export_word(chats, output_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    
    title = doc.add_heading("豆包聊天记录备份", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 共 {len(chats)} 个对话")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    doc.add_heading("目录", level=1)
    for i, chat in enumerate(chats, 1):
        doc.add_paragraph(f"{i}. {chat.get('title', '未命名对话')}")
    doc.add_page_break()
    
    for chat in chats:
        doc.add_heading(chat.get("title", "未命名对话"), level=1)
        for msg in chat.get("messages", []):
            role = msg.get("role", "unknown")
            content = msg.get("content", "")
            role_label = "Assistant" if role == "assistant" else "User" if role == "user" else role
            color = RGBColor(0, 100, 200) if role == "assistant" else RGBColor(0, 150, 0) if role == "user" else RGBColor(128, 128, 128)
            h = doc.add_heading(level=2)
            run = h.add_run(role_label)
            run.font.color.rgb = color
            doc.add_paragraph(content)
        doc.add_page_break()
    
    doc.save(str(output_path))
    print(f"[✓] Word: {output_path}")


# ========== JSON 导出 ==========

def export_json(chats, output_path):
    data = {"export_time": datetime.now(timezone.utc).isoformat(), "count": len(chats), "conversations": chats}
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"[✓] JSON: {output_path}")


# ========== JSONL 导出 ==========

def export_jsonl(chats, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for chat in chats:
            for msg in chat.get("messages", []):
                record = {"chat_id": chat.get("chat_id", ""), "chat_title": chat.get("title", ""), "role": msg.get("role", "unknown"), "content": msg.get("content", "")}
                f.write(json.dumps(record, ensure_ascii=False) + "\n")
    print(f"[✓] JSONL: {output_path}")


# ========== HTML Viewer ==========

def export_html_viewer(chats, output_dir):
    html_dir = output_dir / "html_viewer"
    html_dir.mkdir(exist_ok=True)
    
    sidebar_items = ""
    for i, chat in enumerate(chats):
        title = escape(chat.get("title", "Untitled")[:60])
        msg_count = len(chat.get("messages", []))
        sidebar_items += f'<div class="sidebar-item" onclick="showChat(\'chat_{i}\')">{title} <span class="msg-count">({msg_count})</span></div>\n'
    
    chat_contents = ""
    for i, chat in enumerate(chats):
        title = escape(chat.get("title", "Untitled"))
        messages_html = ""
        for msg in chat.get("messages", []):
            role = msg.get("role", "unknown")
            content = escape(msg.get("content", ""))
            role_class = "assistant" if role == "assistant" else "user"
            role_label = "🤖 Assistant" if role == "assistant" else "👤 User"
            messages_html += f'<div class="message {role_class}"><div class="role">{role_label}</div><div class="content">{content}</div></div>\n'
        chat_contents += f'<div class="chat-content" id="chat_{i}" style="display:none"><h2>{title}</h2>{messages_html}</div>\n'
    
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>豆包 Chat Backup</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, sans-serif; display: flex; height: 100vh; background: #1a1a2e; color: #eee; }}
.sidebar {{ width: 300px; background: #16213e; overflow-y: auto; border-right: 1px solid #0f3460; }}
.sidebar-header {{ padding: 20px; background: #0f3460; text-align: center; }}
.sidebar-item {{ padding: 12px 16px; cursor: pointer; border-bottom: 1px solid #1a1a3e; font-size: 14px; }}
.sidebar-item:hover {{ background: #1a1a3e; }}
.msg-count {{ color: #888; font-size: 12px; }}
.main {{ flex: 1; overflow-y: auto; padding: 20px; }}
.chat-content h2 {{ margin-bottom: 20px; color: #e94560; }}
.message {{ margin-bottom: 16px; padding: 12px 16px; border-radius: 8px; max-width: 80%; }}
.message.user {{ background: #1a3a5c; margin-right: auto; }}
.message.assistant {{ background: #2a1a3e; margin-left: auto; }}
.role {{ font-size: 12px; color: #aaa; margin-bottom: 4px; }}
.content {{ font-size: 14px; line-height: 1.6; white-space: pre-wrap; }}
</style>
</head>
<body>
<div class="sidebar">
<div class="sidebar-header"><h2>豆包 Chat Backup</h2><div style="color:#aaa;font-size:12px;">共 {len(chats)} 个对话</div></div>
{sidebar_items}
</div>
<div class="main">
{chat_contents}
<div id="welcome" style="text-align:center;margin-top:100px;color:#666;"><h2>← 选择一个对话</h2></div>
</div>
<script>
function showChat(id) {{
    document.querySelectorAll('.chat-content').forEach(el => el.style.display = 'none');
    document.getElementById(id).style.display = 'block';
    document.getElementById('welcome').style.display = 'none';
}}
</script>
</body>
</html>'''
    
    with open(html_dir / "index.html", "w", encoding="utf-8") as f:
        f.write(html)
    print(f"[✓] HTML Viewer: {html_dir / 'index.html'}")


# ========== 搜索索引 ==========

def build_search_index(chats, db_path):
    if db_path.exists():
        db_path.unlink()
    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()
    cursor.execute("CREATE VIRTUAL TABLE IF NOT EXISTS chat_fts USING fts5(chat_id, title, content, role, tokenize='unicode61')")
    cursor.execute("CREATE TABLE IF NOT EXISTS chat_meta (chat_id TEXT PRIMARY KEY, title TEXT, message_count INTEGER)")
    for chat in chats:
        chat_id = chat.get("chat_id", "")
        title = chat.get("title", "")
        messages = chat.get("messages", [])
        cursor.execute("INSERT OR REPLACE INTO chat_meta (chat_id, title, message_count) VALUES (?, ?, ?)", (chat_id, title, len(messages)))
        for msg in messages:
            cursor.execute("INSERT INTO chat_fts (chat_id, title, content, role) VALUES (?, ?, ?, ?)", (chat_id, title, msg.get("content", ""), msg.get("role", "unknown")))
    conn.commit()
    conn.close()
    print(f"[✓] Search Index: {db_path}")


def search_chats(db_path, query, limit=20):
    if not db_path.exists():
        return []
    conn = sqlite3.connect(str(db_path))
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT chat_id, title, snippet(chat_fts, 2, '<mark>', '</mark>', '...', 32) FROM chat_fts WHERE chat_fts MATCH ? LIMIT ?", (query, limit))
    results = [{"chat_id": r[0], "title": r[1], "snippet": r[2]} for r in cursor.fetchall()]
    conn.close()
    return results


# ========== ZIP ==========

def create_zip_archive(export_dir, timestamp):
    zip_path = export_dir / f"doubao_backup_{timestamp}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for f in export_dir.iterdir():
            if f.is_file() and f.suffix in ['.md', '.docx', '.pdf', '.json', '.jsonl']:
                zf.write(f, f.name)
            elif f.is_dir() and f.name == 'html_viewer':
                for hf in f.rglob('*'):
                    if hf.is_file():
                        zf.write(hf, f"html_viewer/{hf.name}")
    print(f"[✓] ZIP: {zip_path}")


# ========== 主入口 ==========

FORMAT_HANDLERS = {"markdown": export_markdown, "md": export_markdown, "word": export_word, "docx": export_word, "json": export_json, "jsonl": export_jsonl, "html": export_html_viewer}
FORMAT_EXTENSIONS = {"markdown": ".md", "md": ".md", "word": ".docx", "docx": ".docx", "json": ".json", "jsonl": ".jsonl", "html": ".html"}


def export_all(config=None, formats=None, keyword=None, create_zip=False, build_index=False):
    if config is None:
        config = load_config()
    if formats is None:
        formats = config.get("export_formats", ["markdown", "word"])
    
    backup_dir = Path(os.path.expanduser(config["backup_dir"]))
    chats = load_all_chats(backup_dir)
    if not chats:
        print("[✗] 没有找到已备份的聊天记录")
        return
    
    if keyword:
        chats = [c for c in chats if keyword.lower() in c.get("title", "").lower()]
        print(f"[i] 关键词筛选后: {len(chats)} 个对话")
    
    stats = get_chat_stats(chats)
    print(f"[i] 导出 {stats['total_chats']} 个对话 ({stats['total_messages']} 条消息, {stats['total_words']:,} 字)")
    
    export_dir = backup_dir / "exports"
    export_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    for fmt in formats:
        fmt_lower = fmt.lower()
        handler = FORMAT_HANDLERS.get(fmt_lower)
        ext = FORMAT_EXTENSIONS.get(fmt_lower, ".bin")
        if handler:
            if fmt_lower == "html":
                handler(chats, export_dir)
            else:
                output_path = export_dir / f"doubao_backup_{timestamp}{ext}"
                handler(chats, output_path)
    
    if build_index:
        db_path = export_dir / f"doubao_search_{timestamp}.db"
        build_search_index(chats, db_path)
    
    if create_zip:
        create_zip_archive(export_dir, timestamp)
    
    print(f"\n[✓] 导出完成！目录: {export_dir}")


def main():
    import argparse
    parser = argparse.ArgumentParser(description="豆包聊天记录导出工具")
    parser.add_argument("--format", "-f", nargs="+", choices=["markdown", "md", "word", "docx", "json", "jsonl", "html", "all"], default=["all"])
    parser.add_argument("--keyword", "-k", help="关键词筛选")
    parser.add_argument("--list", "-l", action="store_true", help="列出对话")
    parser.add_argument("--stats", "-s", action="store_true", help="统计信息")
    parser.add_argument("--search", help="全文搜索")
    parser.add_argument("--zip", "-z", action="store_true", help="ZIP归档")
    parser.add_argument("--build-index", action="store_true", help="构建搜索索引")
    args = parser.parse_args()
    
    config = load_config()
    
    if args.search:
        backup_dir = Path(os.path.expanduser(config["backup_dir"]))
        db_files = list(backup_dir.glob("exports/*.db"))
        if not db_files:
            print("[✗] 未找到搜索索引")
            return
        results = search_chats(max(db_files, key=os.path.getmtime), args.search)
        print(f"\n搜索 '{args.search}' 结果 ({len(results)} 条):")
        for r in results:
            print(f"  [{r['title']}] {r['snippet'][:80]}...")
        return
    
    if args.list:
        chats = load_all_chats(Path(os.path.expanduser(config["backup_dir"])))
        if args.keyword:
            chats = [c for c in chats if args.keyword.lower() in c.get("title", "").lower()]
        print(f"\n{'ID':<45} {'标题'}")
        print("-" * 80)
        for c in chats:
            print(f"{c.get('chat_id', '')[:43]:<45} {c.get('title', '')[:35]}")
        print(f"\n共 {len(chats)} 个对话")
        return
    
    if args.stats:
        chats = load_all_chats(Path(os.path.expanduser(config["backup_dir"])))
        stats = get_chat_stats(chats)
        print(f"\n{'='*50}")
        print(f"豆包聊天记录统计")
        print(f"{'='*50}")
        print(f"  对话总数: {stats['total_chats']}")
        print(f"  消息总数: {stats['total_messages']}")
        print(f"  总字数:   {stats['total_words']:,}")
        print(f"{'='*50}")
        return
    
    formats = []
    for f in args.format:
        if f == "all":
            formats = ["markdown", "word", "json", "jsonl", "html"]
            break
        formats.append(f)
    
    export_all(config=config, formats=formats, keyword=args.keyword, create_zip=args.zip, build_index=args.build_index)


if __name__ == "__main__":
    main()
